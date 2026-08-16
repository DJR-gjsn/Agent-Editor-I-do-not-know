"""
MCP PDF 文档服务模块
提供 PDF 读取、创建、合并工具
"""

import os
import tempfile
import time
from pathlib import Path

from flask import jsonify, request, send_file

from . import tool_registry

# ============================================================
# 输出目录
# ============================================================
_OUT_DIR = Path(tempfile.gettempdir()) / "mcp_pdf_output"
_OUT_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================
# 工具定义
# ============================================================
PDF_READ_DEF = {
    "name": "pdf_read",
    "description": (
        "读取 PDF 文件并提取纯文本内容。"
        "可以读取全部页面或指定页码范围。"
        "适合需要从 PDF 文档中提取信息、阅读 PDF 内容时使用。"
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "PDF 文件的完整路径"},
            "start_page": {"type": "integer", "description": "起始页码（从 1 开始），默认 1"},
            "end_page": {"type": "integer", "description": "结束页码（包含），不指定则到最后一页"},
        },
        "required": ["path"],
    },
}

PDF_CREATE_DEF = {
    "name": "pdf_create",
    "description": (
        "从文本/Markdown 内容创建 PDF 文件，或从已有文件（.txt/.md/.json/.csv）转换为 PDF。"
        "适合生成报告、导出文档为 PDF 时使用。"
        "两种用法：1) 传入 content + title 创建新 PDF；2) 传入 source_file 把已有文件转成 PDF。"
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "content": {"type": "string", "description": "要写入 PDF 的文本内容。支持 Markdown 格式。与 source_file 二选一。"},
            "title": {"type": "string", "description": "PDF 文档标题，也用作文件名（不含扩展名）"},
            "source_file": {"type": "string", "description": "已有文本文件的路径，将其内容转换为 PDF。支持 .txt/.md/.json/.csv。与 content 二选一。"},
        },
        "required": ["title"],
    },
}

PDF_MERGE_DEF = {
    "name": "pdf_merge",
    "description": "合并多个 PDF 文件为一个。按参数顺序合并。",
    "parameters": {
        "type": "object",
        "properties": {
            "files": {
                "type": "array",
                "items": {"type": "string"},
                "description": "要合并的 PDF 文件路径列表，如 ['a.pdf', 'b.pdf']",
            },
            "output_name": {"type": "string", "description": "合并后的输出文件名（不含扩展名），默认 'merged'", "default": "merged"},
        },
        "required": ["files"],
    },
}


# ============================================================
# 执行器
# ============================================================
def _exec_pdf_read(args: dict) -> str:
    path = (args.get("path") or "").strip()
    start = int(args.get("start_page", 1))
    end = int(args.get("end_page", 0)) or None

    p = Path(path).expanduser()
    if not p.exists():
        return f"❌ 文件不存在: {path}"

    try:
        from pypdf import PdfReader
        reader = PdfReader(str(p))
        total = len(reader.pages)

        if end is None:
            end = total
        start = max(1, start)
        end = min(end, total)

        lines = [f"📄 {p.name} (共 {total} 页，提取 {start}-{end} 页):"]
        for i in range(start - 1, end):
            page = reader.pages[i]
            text = page.extract_text() or "(无文本)"
            lines.append(f"\n--- 第 {i + 1} 页 ---\n{text[:2000]}")

        result = "\n".join(lines)
        if len(result) > 5000:
            result = result[:5000] + "\n\n... [截断，原文更长]"
        return result
    except ImportError:
        return "❌ 需要安装 pypdf: pip install pypdf"
    except Exception as e:
        return f"❌ PDF 读取失败: {str(e)}"


def _exec_pdf_create(args: dict) -> str:
    content = args.get("content", "")
    title = args.get("title", "document")
    source_file = args.get("source_file", "")

    # 如果指定了源文件，读取文件内容
    if source_file and not content:
        sp = Path(source_file).expanduser()
        if not sp.exists():
            return f"❌ 源文件不存在: {source_file}"
        if not title or title == "document":
            title = sp.stem
        try:
            ext = sp.suffix.lower()
            if ext == ".docx":
                # 从 Word 文档提取文本
                try:
                    from docx import Document
                    doc = Document(str(sp))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    # 也提取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text for cell in row.cells)
                            if row_text.strip():
                                paragraphs.append(row_text)
                    content = "\n\n".join(paragraphs)
                except ImportError:
                    return "❌ 需要安装 python-docx 来读取 Word 文件: pip install python-docx"
            elif ext in (".txt", ".md", ".json", ".csv", ".py", ".js", ".html", ".css", ".log"):
                content = sp.read_text(encoding="utf-8", errors="replace")
            else:
                # 尝试作为文本读取
                try:
                    content = sp.read_text(encoding="utf-8", errors="replace")
                except Exception:
                    return f"❌ 不支持的文件格式: {ext}。支持 .txt/.md/.docx/.json/.csv"
        except Exception as e:
            return f"❌ 读取源文件失败: {e}"

    if not content:
        return "错误: 请提供 content 文本内容 或 source_file 文件路径"

    out_path = _OUT_DIR / f"{title}.pdf"
    if out_path.exists():
        out_path = _OUT_DIR / f"{title}_{int(time.time())}.pdf"

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch

        doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                                rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()
        story = []

        # 标题
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 12))

        # 将内容按段落拆分
        for para in content.split("\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(para, styles["Normal"]))
                story.append(Spacer(1, 6))

        doc.build(story)
        return f"✅ PDF 已创建: {out_path}\n📏 {out_path.stat().st_size:,} bytes"
    except ImportError:
        # 简易回退：纯文本 PDF
        return _create_simple_pdf(content, title, out_path)
    except Exception as e:
        return f"❌ PDF 创建失败: {str(e)}"


def _find_cjk_font() -> str:
    """Find a CJK-capable TTF font on the system. Returns path or empty string."""
    import os as _os
    font_dir = None
    for d in ["C:/Windows/Fonts", "/usr/share/fonts", "/System/Library/Fonts"]:
        if _os.path.isdir(d):
            font_dir = d
            break
    if not font_dir:
        return ""
    candidates = ["Noto Sans SC.ttf", "simhei.ttf", "simsunb.ttf",
                   "NotoSansSC-VF.ttf", "SimsunExtG.ttf"]
    for name in candidates:
        path = _os.path.join(font_dir, name)
        if _os.path.isfile(path):
            return path
    return ""


def _create_simple_pdf(content: str, title: str, out_path: Path) -> str:
    """纯文本 PDF 创建（无需 reportlab），支持 CJK 字符"""
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        # Try CJK font for Unicode support, fallback to Arial
        cjk_font = _find_cjk_font()
        if cjk_font:
            pdf.add_font("CJK", "", cjk_font, uni=True)
            pdf.set_font("CJK", size=12)
        else:
            pdf.set_font("Arial", size=12)
        for line in content.split("\n"):
            pdf.cell(200, 10, txt=line[:120], ln=True)
        pdf.output(str(out_path))
        return f"✅ PDF 已创建 (fpdf): {out_path}\n💡 安装 reportlab 获得更好排版: pip install reportlab"
    except ImportError:
        return f"❌ 需要安装 PDF 库。请运行:\npip install fpdf  (基础)\n或\npip install reportlab  (推荐)"


def _exec_pdf_merge(args: dict) -> str:
    files = args.get("files", [])
    output_name = args.get("output_name", "merged")

    if not files:
        return "错误: files 不能为空"

    out_path = _OUT_DIR / f"{output_name}.pdf"

    try:
        from pypdf import PdfReader, PdfWriter
        writer = PdfWriter()
        for f in files:
            p = Path(f).expanduser()
            if not p.exists():
                return f"❌ 文件不存在: {f}"
            reader = PdfReader(str(p))
            for page in reader.pages:
                writer.add_page(page)

        writer.write(str(out_path))
        return f"✅ {len(files)} 个 PDF 已合并: {out_path}\n📏 {out_path.stat().st_size:,} bytes"
    except ImportError:
        return "❌ 需要安装 pypdf: pip install pypdf"
    except Exception as e:
        return f"❌ PDF 合并失败: {str(e)}"


_tool_list = [
    (PDF_READ_DEF, _exec_pdf_read),
    (PDF_CREATE_DEF, _exec_pdf_create),
    (PDF_MERGE_DEF, _exec_pdf_merge),
]

for tool_def, executor in _tool_list:
    tool_registry.register(tool_def["name"], tool_def, executor)


# ============================================================
# Flask 路由
# ============================================================
def register_routes(app, http_session=None):
    @app.route("/api/pdf/test-connection", methods=["POST"])
    def pdf_test_connection():
        """测试 PDF 模块可用性"""
        checks = {}
        for mod in ["pypdf", "reportlab", "fpdf"]:
            try:
                __import__(mod)
                checks[mod] = "可用"
            except ImportError:
                checks[mod] = "未安装"

        available = [m for m, s in checks.items() if s == "可用"]
        return jsonify({
            "success": len(available) > 0,
            "message": f"PDF 模块: {', '.join(available)} 可用" if available else "PDF 模块未安装。pip install pypdf fpdf",
            "modules": checks,
            "output_dir": str(_OUT_DIR),
        })

    @app.route("/api/pdf/download/<filename>")
    def pdf_download(filename):
        """下载生成的 PDF"""
        filepath = _OUT_DIR / f"{filename}.pdf"
        if filepath.exists():
            return send_file(str(filepath), mimetype="application/pdf", as_attachment=True)
        return jsonify({"error": "文件不存在"}), 404
