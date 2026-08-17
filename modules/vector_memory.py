"""
Vector Memory 向量记忆模块
embeddings_index: 将文档向量化并存入本地向量库
embeddings_search: 语义搜索，余弦相似度匹配 top-K
"""
import json
import math
import os
import threading
import time

import requests as _requests
from flask import jsonify, request

from . import tool_registry
from .config import get_config
from .utils import get_request_api_config

# ============================================================
# 向量存储（线程安全 + 落盘持久化）
# ============================================================
_lock = threading.RLock()
_store = []  # [{"id": str, "text": str, "embedding": list[float], "created_at": str}]
_next_id = 0

# 存储路径覆盖（init_vector_store 设置；测试/多实例用）。None 时回退到
# 环境变量 VECTOR_STORE_PATH，再回退到默认 data/vector_store/vector_store.json
_STORE_PATH_OVERRIDE = None

# 默认存储文件：项目根/data/vector_store/vector_store.json
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_DEFAULT_STORE_PATH = os.path.join(_PROJECT_ROOT, "data", "vector_store", "vector_store.json")


def _store_path() -> str:
    """当前向量库存储文件路径（override > 环境变量 > 默认）"""
    if _STORE_PATH_OVERRIDE:
        return _STORE_PATH_OVERRIDE
    env = os.environ.get("VECTOR_STORE_PATH")
    if env:
        return env
    return _DEFAULT_STORE_PATH


def _save_to_disk():
    """将内存向量库原子写入磁盘（tmp + os.replace，避免写一半损坏）"""
    with _lock:
        data = {"version": 1, "next_id": _next_id, "documents": _store}
        path = _store_path()
        tmp = f"{path}.tmp"
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(tmp, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False)
            os.replace(tmp, path)
        except OSError:
            # 落盘失败不阻塞主流程（向量库仍在内存中可用）
            try:
                if os.path.exists(tmp):
                    os.remove(tmp)
            except OSError:
                pass


def _load_from_disk():
    """从磁盘加载向量库（幂等）。文件缺失/损坏时以空库启动，不崩溃。"""
    global _store, _next_id
    path = _store_path()
    with _lock:
        _store = []
        _next_id = 0
        if not os.path.exists(path):
            return
        try:
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
        except Exception:
            return  # 损坏文件：空库启动
        docs = data.get("documents", []) if isinstance(data, dict) else []
        valid = [
            d for d in docs
            if isinstance(d, dict) and d.get("id") and d.get("text") is not None
            and isinstance(d.get("embedding"), list)
        ]
        _store = valid
        # 恢复 next_id：取 文件记录的 next_id、已有 id 最大值+1、文档数 三者的最大值，
        # 保证新文档 id 不会与旧文档重叠
        max_id = -1
        for d in _store:
            iid = str(d["id"])
            if iid.startswith("vec_"):
                try:
                    max_id = max(max_id, int(iid[4:]))
                except ValueError:
                    pass
        nid = data.get("next_id", 0) if isinstance(data, dict) else 0
        if not isinstance(nid, int) or nid < 0:
            nid = 0
        _next_id = max(nid, max_id + 1, len(_store))


def init_vector_store(path=None):
    """重定向向量库存储路径并重新从盘加载。

    - 传路径：后续读写都落到该文件（测试/多实例隔离）
    - 传 None：回到默认路径（环境变量或项目默认位置）
    """
    global _STORE_PATH_OVERRIDE
    _STORE_PATH_OVERRIDE = path
    _load_from_disk()


def _clear_store():
    """清空向量库（内存 + 磁盘）"""
    global _store, _next_id
    with _lock:
        _store = []
        _next_id = 0
    _save_to_disk()


def _get_stats() -> dict:
    """获取向量库统计"""
    with _lock:
        dims = len(_store[0]["embedding"]) if _store else 0
        return {
            "doc_count": len(_store),
            "dimensions": dims,
            "total_chars": sum(len(d["text"]) for d in _store),
        }


# ============================================================
# Embedding API 调用
# ============================================================
def _get_embedding(text: str, api_base: str = None, api_key: str = None) -> list:
    """获取文本向量。优先用 API embeddings，不可用时回退到本地 TF-IDF"""
    cfg = get_config()
    req_cfg = get_request_api_config()
    base = api_base or req_cfg.get("api_base") or cfg["api_base"]
    key = api_key or req_cfg.get("api_key") or cfg["api_key"]

    # 未配置有效 key（空或默认占位符）时直接本地向量化，避免无谓的网络等待
    if not key or "your-api-key" in key:
        return _tfidf_embed(text)

    # 尝试 API embeddings（仅首次，失败后跳过）
    global _api_available
    if _api_available:
        try:
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {key}",
            }
            payload = {"model": "text-embedding-3-small", "input": text}
            resp = _requests.post(
                f"{base}/embeddings", headers=headers, json=payload, timeout=2,
            )
            if resp.status_code == 200:
                body = resp.json()
                return body["data"][0]["embedding"]
        except Exception:
            _api_available = False  # 失败后不再尝试

    # 回退：本地 TF-IDF 向量化
    return _tfidf_embed(text)


# ============================================================
# TF-IDF 本地向量化（无需 API，无需网络）
# ============================================================
_tfidf_vectorizer = None
_api_available = True  # 缓存 API 可用状态，失败一次后跳过


def _tfidf_embed(text: str) -> list:
    """TF-IDF 向量化，使用字符 n-gram 适配中文"""
    global _tfidf_vectorizer

    from sklearn.feature_extraction.text import TfidfVectorizer

    with _lock:
        corpus = [d["text"] for d in _store]
        corpus.append(text)

        try:
            # 字符级 n-gram (2-4) 对中文和英文混合文本都有效
            _tfidf_vectorizer = TfidfVectorizer(
                analyzer='char',
                ngram_range=(2, 4),
                max_features=512,
            )
            tfidf_matrix = _tfidf_vectorizer.fit_transform(corpus)
            vec = tfidf_matrix[-1].toarray()[0].tolist()

            # 更新所有已存储文档的向量
            if _store:
                for i, doc in enumerate(_store):
                    doc["embedding"] = tfidf_matrix[i].toarray()[0].tolist()

            return vec
        except Exception:
            return _fallback_hash(text, 512)


def _fallback_hash(text: str, dims: int = 256) -> list:
    """极端回退：字符级简单向量（每个维度基于字符哈希）"""
    import hashlib
    vec = [0.0] * dims
    chars = list(text)
    for i, ch in enumerate(chars):
        h = int(hashlib.md5(ch.encode()).hexdigest()[:8], 16)
        idx = (h + i * 31) % dims
        vec[idx] += 1.0
    # 归一化
    norm = math.sqrt(sum(v * v for v in vec))
    if norm > 0:
        vec = [v / norm for v in vec]
    return vec


# ============================================================
# 余弦相似度
# ============================================================
def _cosine_similarity(a: list, b: list) -> float:
    """两个向量的余弦相似度 (0~1, 越大越相似)"""
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x * x for x in a))
    norm_b = math.sqrt(sum(y * y for y in b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


# ============================================================
# 工具定义
# ============================================================
INDEX_DEF = {
    "name": "embeddings_index",
    "description": (
        "将文档添加到向量知识库。传入文本内容，系统会自动向量化并存储。"
        "适合将重要信息、参考资料、对话摘要等存入知识库，供后续语义搜索使用。"
        "多次调用可积累多条文档。"
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "text": {
                "type": "string",
                "description": "要添加到向量知识库的文档内容",
            },
        },
        "required": ["text"],
    },
}

SEARCH_DEF = {
    "name": "embeddings_search",
    "description": (
        "在向量知识库中进行语义搜索。根据查询内容，找到知识库中意思最相近的文档。"
        "不是关键词匹配，而是理解语义后找相关内容。"
        "适合查找之前存储的参考资料、对话记录、专业知识等。"
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "query": {
                "type": "string",
                "description": "搜索查询，用自然语言描述你想找什么",
            },
            "top_k": {
                "type": "integer",
                "description": "返回最相似的前几条结果，默认 3，最大 10",
                "default": 3,
            },
        },
        "required": ["query"],
    },
}


# ============================================================
# 工具执行器
# ============================================================
def _exec_index(args: dict) -> str:
    """AI 可调用：将文档加入向量库"""
    text = (args.get("text") or "").strip()
    if not text:
        return "错误: text 不能为空"
    if len(text) > 8000:
        return "错误: 文本过长（最多 8000 字符）"

    global _next_id
    try:
        embedding = _get_embedding(text)
    except Exception as e:
        return f"❌ 向量化失败: {e}"

    with _lock:
        doc_id = f"vec_{_next_id}"
        _next_id += 1
        _store.append({
            "id": doc_id,
            "text": text,
            "embedding": embedding,
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        })
        total = len(_store)
    _save_to_disk()  # 持久化：写入磁盘，重启后不丢失

    preview = text[:100] + ("..." if len(text) > 100 else "")
    return f"✅ 已添加到向量知识库 (ID: {doc_id})\n当前共 {total} 条文档\n内容: {preview}"


def _exec_search(args: dict) -> str:
    """AI 可调用：语义搜索向量库"""
    query = (args.get("query") or "").strip()
    top_k = min(int(args.get("top_k", 3)), 10)

    if not query:
        return "错误: query 不能为空"

    with _lock:
        if not _store:
            return "📭 向量知识库为空。请先用 embeddings_index 添加文档。"

    # 向量化查询（在锁外执行，避免长时间持锁）
    try:
        query_emb = _get_embedding(query)
    except Exception as e:
        return f"❌ 查询向量化失败: {e}"

    # 计算相似度并排序
    with _lock:
        scored = []
        for doc in _store:
            score = _cosine_similarity(query_emb, doc["embedding"])
            scored.append((score, doc))

        scored.sort(key=lambda x: x[0], reverse=True)
        top = scored[:top_k]

    lines = [f"🔍 语义搜索 '{query[:60]}' 的结果 (top {len(top)}):"]
    for i, (score, doc) in enumerate(top, 1):
        pct = f"{score * 100:.0f}%"
        text_preview = doc["text"][:200]
        lines.append(
            f"\n[{i}] 相似度: {pct} | {doc['created_at']}\n"
            f"    {text_preview}"
        )
        if len(doc["text"]) > 200:
            lines.append("    ...")

    return "\n".join(lines)


# ============================================================
# 注册工具
# ============================================================
tool_registry.register("embeddings_index", INDEX_DEF, _exec_index)
tool_registry.register("embeddings_search", SEARCH_DEF, _exec_search)


# ============================================================
# Flask 路由
# ============================================================
def register_routes(app, http_session=None):
    @app.route("/api/vector-memory/documents", methods=["GET", "POST"])
    def vm_documents():
        """列出文档 或 添加文档"""
        if request.method == "POST":
            data = request.get_json(force=True)
            text = (data.get("text") or "").strip()
            if not text:
                return jsonify({"success": False, "error": "text 不能为空"})
            result = _exec_index({"text": text})
            return jsonify({"success": True, "result": result})
        # GET：列出向量库中的所有文档（不含向量数据）
        with _lock:
            docs = [{
                "id": d["id"],
                "text": d["text"][:200],
                "created_at": d["created_at"],
                "dimensions": len(d["embedding"]),
            } for d in _store]
        return jsonify({"success": True, "documents": docs, "count": len(docs)})

    @app.route("/api/vector-memory/documents", methods=["DELETE"])
    def vm_clear_docs():
        """清空向量库"""
        _clear_store()
        return jsonify({"success": True, "message": "向量库已清空"})

    @app.route("/api/vector-memory/import-file", methods=["POST"])
    def vm_import_file():
        """从上传文件提取文本并导入向量知识库（支持 txt/md/csv/json/pdf/docx/xlsx）"""
        f = request.files.get("file")
        if not f or not f.filename:
            return jsonify({"success": False, "error": "未收到文件"})
        name = f.filename or ""
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        raw = f.read()

        import io

        text = ""
        try:
            if ext in ("txt", "md", "csv", "json"):
                text = raw.decode("utf-8", errors="replace")
            elif ext == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(raw))
                text = "\n".join((p.extract_text() or "") for p in reader.pages)
            elif ext == "docx":
                from docx import Document
                doc = Document(io.BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs)
                for tbl in doc.tables:
                    for row in tbl.rows:
                        text += "\n" + "\t".join(c.text for c in row.cells)
            elif ext == "xlsx":
                from openpyxl import load_workbook
                wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
                parts = []
                for ws in wb.worksheets:
                    for row in ws.iter_rows(values_only=True):
                        vals = [str(c) for c in row if c is not None]
                        if vals:
                            parts.append("\t".join(vals))
                text = "\n".join(parts)
            elif ext == "doc":
                return jsonify({"success": False, "error": "暂不支持旧版 .doc，请另存为 .docx 或 .txt"})
            else:
                text = raw.decode("utf-8", errors="replace")
        except Exception as e:
            return jsonify({"success": False, "error": f"解析 .{ext} 失败: {e}"})

        text = (text or "").strip()
        if not text:
            return jsonify({"success": False, "error": "未能从文件中提取到文本内容"})

        result = _exec_index({"text": text[:20000]})
        return jsonify({"success": True, "result": result, "chars": len(text),
                        "filename": name, "ext": ext})

    @app.route("/api/vector-memory/search", methods=["POST"])
    def vm_search():
        """直接搜索向量库（非 AI 调用）"""
        data = request.get_json(force=True)
        query = (data.get("query") or "").strip()
        top_k = min(int(data.get("top_k", 3)), 10)

        if not query:
            return jsonify({"success": False, "error": "query 不能为空"})

        result = _exec_search({"query": query, "top_k": top_k})
        return jsonify({"success": True, "result": result})

    @app.route("/api/vector-memory/stats", methods=["GET"])
    def vm_stats():
        return jsonify({"success": True, "stats": _get_stats()})


# ============================================================
# 启动时从磁盘加载（持久化数据恢复）
# ============================================================
_load_from_disk()
